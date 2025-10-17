"""
FSC Word Document Generator - Self-contained
No dependencies on FSC Developer plugin
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional
from datetime import datetime


class FSCContentValidator:
    """
    Validates FSC content dictionary against expected contract.
    Self-contained - no imports from other plugins.
    """
    
    SCHEMA_VERSION = "1.0"
    
    @staticmethod
    def validate(content: Dict) -> tuple[bool, List[str], List[str]]:
        """
        Validate FSC content structure.
        
        Returns:
            (is_valid, errors, warnings)
        """
        errors = []
        warnings = []
        
        # Check schema version
        metadata = content.get('metadata', {})
        schema_version = metadata.get('schema_version', 'unknown')
        if schema_version != FSCContentValidator.SCHEMA_VERSION:
            warnings.append(f"Schema version mismatch: expected {FSCContentValidator.SCHEMA_VERSION}, got {schema_version}")
        
        # Required keys
        required = [
            'system_name', 'introduction', 'safety_goal_summary',
            'functional_safety_requirements', 'safety_mechanisms',
            'architectural_allocation', 'verification_strategy'
        ]
        
        for key in required:
            if key not in content:
                errors.append(f"Missing required key: {key}")
        
        # Validate FSRs
        fsrs = content.get('functional_safety_requirements', [])
        if not isinstance(fsrs, list):
            errors.append("functional_safety_requirements must be a list")
        elif len(fsrs) == 0:
            warnings.append("No FSRs found")
        else:
            for i, fsr in enumerate(fsrs):
                if not isinstance(fsr, dict):
                    errors.append(f"FSR {i} is not a dictionary")
                    continue
                
                # Check essential FSR fields
                if 'id' not in fsr:
                    errors.append(f"FSR {i} missing 'id'")
                if 'description' not in fsr:
                    errors.append(f"FSR {i} missing 'description'")
        
        # Validate Safety Mechanisms
        sms = content.get('safety_mechanisms', [])
        if not isinstance(sms, list):
            errors.append("safety_mechanisms must be a list")
        elif len(sms) == 0:
            warnings.append("No safety mechanisms found")
        
        is_valid = len(errors) == 0
        return (is_valid, errors, warnings)


class FSCWordGenerator:
    """
    Generate ISO 26262-3:2018 Clause 7 compliant Word documents.
    Reads structured content from working memory.
    """
    
    def __init__(self):
        self.validator = FSCContentValidator()
    
    def generate(self, system_name=None, goals_data=None, fsrs_data=None, 
                 strategies_data=None, allocation_data=None, structured_content=None):
        """
        Generate FSC Word document.
        
        Args:
            structured_content: Dictionary from fsc_structured_content (contract v1.0)
            other args: Legacy support
        """
        
        # Prefer structured content
        if structured_content:
            return self._generate_from_structured(structured_content)
        else:
            # Legacy path
            return self._generate_legacy(system_name, goals_data, fsrs_data, 
                                         strategies_data, allocation_data)
    
    def _generate_from_structured(self, content: Dict) -> Document:
        """
        Generate document from structured content dictionary.
        No imports needed - just reads the dictionary.
        """
        
        # Validate
        is_valid, errors, warnings = self.validator.validate(content)
        
        if not is_valid:
            raise ValueError(f"Invalid FSC content: {'; '.join(errors)}")
        
        if warnings:
            print(f"⚠️ Warnings: {'; '.join(warnings)}")
        
        # Create document
        doc = Document()
        
        # Extract data
        system_name = content.get('system_name', 'Unknown System')
        introduction = content.get('introduction', '')
        safety_goal_summary = content.get('safety_goal_summary', '')
        fsrs = content.get('functional_safety_requirements', [])
        sms = content.get('safety_mechanisms', [])
        allocation = content.get('architectural_allocation', '')
        verification = content.get('verification_strategy', '')
        
        # Title Page
        self._add_title_page(doc, system_name)
        
        # Table of Contents placeholder
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph('[Auto-generated TOC would go here]')
        doc.add_page_break()
        
        # 1. Introduction
        doc.add_heading('1. Introduction', level=1)
        doc.add_paragraph(introduction)
        doc.add_paragraph()
        
        # 2. Safety Goals
        doc.add_heading('2. Safety Goals', level=1)
        doc.add_paragraph(safety_goal_summary)
        doc.add_paragraph()
        
        # 3. Functional Safety Requirements
        doc.add_heading('3. Functional Safety Requirements', level=1)
        doc.add_paragraph(
            'The following functional safety requirements have been derived '
            'from the safety goals per ISO 26262-3:2018, Clause 7.4.2.'
        )
        doc.add_paragraph()
        
        for idx, fsr in enumerate(fsrs, 1):
            self._add_fsr_section(doc, idx, fsr)
        
        # 4. Safety Mechanisms
        doc.add_heading('4. Safety Mechanisms', level=1)
        doc.add_paragraph(
            'The following safety mechanisms implement the functional safety requirements '
            'per ISO 26262-3:2018, Clause 7.4.2.3.'
        )
        doc.add_paragraph()
        
        self._add_safety_mechanisms_section(doc, sms)
        
        # 5. Traceability Matrix
        doc.add_heading('5. FSR-SM Traceability Matrix', level=1)
        self._add_traceability_matrix(doc, fsrs, sms)
        
        # 6. Architectural Allocation
        doc.add_heading('6. Architectural Allocation', level=1)
        doc.add_paragraph(allocation)
        doc.add_paragraph()
        
        # 7. Verification and Validation
        doc.add_heading('7. Verification and Validation', level=1)
        doc.add_paragraph(verification)
        doc.add_paragraph()
        
        # 8. Document Information
        self._add_document_info(doc, content.get('metadata', {}))
        
        return doc
    
    def _add_title_page(self, doc: Document, system_name: str):
        """Add title page"""
        title = doc.add_heading('Functional Safety Concept', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # System name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\n\n{system_name}\n\n')
        run.font.size = Pt(18)
        run.font.bold = True
        
        # Compliance statement
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('ISO 26262-3:2018, Clause 7\nFunctional Safety Concept')
        run.font.size = Pt(12)
        run.italic = True
        
        doc.add_page_break()
    
    def _add_fsr_section(self, doc: Document, idx: int, fsr: Dict):
        """Add single FSR section"""
        fsr_id = fsr.get('id', f'FSR-{idx:03d}')
        
        doc.add_heading(f'3.{idx} {fsr_id}', level=2)
        
        # FSR details table
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Populate table
        table.cell(0, 0).text = 'Requirement ID'
        table.cell(0, 1).text = fsr_id
        
        table.cell(1, 0).text = 'Description'
        table.cell(1, 1).text = fsr.get('description', 'N/A')
        
        table.cell(2, 0).text = 'Type'
        table.cell(2, 1).text = fsr.get('type', 'N/A').title()
        
        table.cell(3, 0).text = 'ASIL'
        table.cell(3, 1).text = fsr.get('asil', 'QM')
        
        table.cell(4, 0).text = 'Safety Goal'
        table.cell(4, 1).text = fsr.get('safety_goal_id', 'N/A')
        
        table.cell(5, 0).text = 'Safe State'
        table.cell(5, 1).text = fsr.get('safe_state', 'TBD')
        
        table.cell(6, 0).text = 'FTTI'
        table.cell(6, 1).text = fsr.get('ftti', 'TBD')
        
        table.cell(7, 0).text = 'Verification Method'
        table.cell(7, 1).text = fsr.get('verification_method', 'TBD')
        
        # Validation criteria
        doc.add_paragraph()
        doc.add_paragraph('Validation Criteria:', style='Heading 3')
        
        criteria = fsr.get('validation_criteria', [])
        if criteria:
            for criterion in criteria:
                doc.add_paragraph(criterion, style='List Bullet')
        else:
            doc.add_paragraph('To be defined', style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_safety_mechanisms_section(self, doc: Document, sms: List[Dict]):
        """Add safety mechanisms, grouped by type"""
        from collections import defaultdict
        
        # Group by type
        grouped = defaultdict(list)
        for sm in sms:
            sm_type = sm.get('type', 'other')
            grouped[sm_type].append(sm)
        
        subsection = 1
        for sm_type in ['detection', 'mitigation', 'control', 'other']:
            mechanisms = grouped.get(sm_type, [])
            if not mechanisms:
                continue
            
            doc.add_heading(f'4.{subsection} {sm_type.title()} Mechanisms', level=2)
            
            for sm in mechanisms:
                p = doc.add_paragraph(style='List Bullet')
                
                # ID and name
                sm_id = sm.get('id', 'SM-?')
                sm_name = sm.get('name', sm_id)
                run = p.add_run(f'{sm_id} - {sm_name}: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                
                # Description
                p.add_run(sm.get('description', 'N/A') + ' ')
                
                # Coverage
                coverage = sm.get('fsr_coverage', [])
                if coverage:
                    run = p.add_run(f"(Covers: {', '.join(coverage)})")
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)
            
            subsection += 1
            doc.add_paragraph()
    
    def _add_traceability_matrix(self, doc: Document, fsrs: List[Dict], sms: List[Dict]):
        """Add FSR-SM traceability matrix"""
        if not fsrs or not sms:
            doc.add_paragraph('Insufficient data for traceability matrix.')
            return
        
        # Create table
        num_rows = len(fsrs) + 1
        num_cols = len(sms) + 1
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        table.cell(0, 0).text = 'FSR \\ SM'
        for j, sm in enumerate(sms, 1):
            table.cell(0, j).text = sm.get('id', f'SM-{j}')
        
        # Fill matrix
        for i, fsr in enumerate(fsrs, 1):
            fsr_id = fsr.get('id', f'FSR-{i}')
            table.cell(i, 0).text = fsr_id
            
            for j, sm in enumerate(sms, 1):
                coverage = sm.get('fsr_coverage', [])
                if fsr_id in coverage:
                    cell = table.cell(i, j)
                    cell.text = '✓'
                    # Center align
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def _add_document_info(self, doc: Document, metadata: Dict):
        """Add document metadata section"""
        doc.add_heading('8. Document Information', level=1)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Generation Date'
        table.cell(0, 1).text = metadata.get('generation_date', 'N/A')
        
        table.cell(1, 0).text = 'Generator'
        table.cell(1, 1).text = metadata.get('generator_plugin', 'N/A')
        
        table.cell(2, 0).text = 'Schema Version'
        table.cell(2, 1).text = metadata.get('schema_version', 'N/A')
        
        table.cell(3, 0).text = 'Document Version'
        table.cell(3, 1).text = '1.0 (Draft)'
    
    def _generate_legacy(self, system_name, goals_data, fsrs_data, 
                        strategies_data, allocation_data):
        """Legacy generation - keep your existing implementation"""
        # Your current code here
        pass