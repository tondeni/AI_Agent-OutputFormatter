# utils.py - Common utilities for document formatting
import os
import json
import re
from datetime import datetime
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from cat.log import log

def detect_document_type(content, working_memory):
    """
    Detect document type from content and working memory.
    PRIORITY: Check working memory first, then content markers.
    
    Args:
        content (str): Document content
        working_memory (dict): Cat's working memory
        
    Returns:
        str: Document type identifier or None
    """
    
    # CRITICAL FIX: Check working memory FIRST - this is the most reliable
    doc_type = working_memory.get("document_type")
    
    if doc_type:
        log.info(f"📌 Document type from working memory: {doc_type}")
        
        # Additional check for HARA review to avoid confusion
        review_type = working_memory.get("review_type")
        if review_type == "HARA_COMPLIANCE":
            log.info("📌 Confirmed HARA review via review_type marker")
            return "hara_review"
        
        return doc_type
    
    # FALLBACK: Check content patterns if working memory doesn't have explicit type
    content_lower = content.lower()
    
    # Check for HARA REVIEW first (most specific)
    if any(marker in content for marker in [
        "# HARA COMPLIANCE REVIEW REPORT",
        "HARA Review per ISO 26262-3",
        "HARA_COMPLIANCE"
    ]):
        log.info("📌 Detected HARA review from content markers")
        return "hara_review"
    
    # Check for HARA-specific patterns in content
    if "hara" in content_lower and any(marker in content_lower for marker in [
        "hazard analysis and risk assessment",
        "asil determination",
        "severity", "exposure", "controllability",
        "hazardous event"
    ]):
        # Check if this is a review or the HARA itself
        if any(review_marker in content_lower for review_marker in [
            "**status:**", "**comment:**", "**hint for improvement:**",
            "pass / fail", "review checklist"
        ]):
            log.info("📌 Detected HARA review from content patterns")
            return "hara_review"
        else:
            log.info("📌 Detected HARA document from content patterns")
            return "hara"
    
    # Check for Item Definition Review (less specific, check after HARA review)
    if "item definition" in content_lower and any(marker in content_lower for marker in [
        "**status:**", "**comment:**", "**hint for improvement:**",
        "iso 26262-3 clause 5"
    ]):
        log.info("📌 Detected item definition review from content patterns")
        return "item_definition_review"
    
    # Check for Item Definition
    if any(marker in content for marker in [
        "# Item Definition:",
        "## Item Definition Document",
        "### Item Description"
    ]):
        log.info("📌 Detected item definition from content markers")
        return "item_definition"
    
    # Check for Safety Goals
    if "safety goal" in content_lower and working_memory.get("hara_stage") == "safety_goals_derived":
        log.info("📌 Detected safety goals from content and stage")
        return "safety_goals"
    
    # Check for FSR document
    if any(marker in content for marker in [
        "Functional Safety Requirements",
        "FSR-",
        "**FSR-",
        "7.4.2.1",
        "7.4.2.2"
        ]) and "derived" in content.lower():
        log.info("📌 Detected FSR document from content markers")
        return "fsr"
    
    # No clear type detected
    log.info("📌 No specific document type detected")
    return None

def parse_review_content(content):
    """
    Parse review content into structured review items.
    
    Args:
        content (str): Raw review content with markdown formatting
        
    Returns:
        list: List of review item dictionaries
    """
    reviews = []
    current_review = {}
    
    # Split by common separators
    sections = re.split(r'\n\s*---\s*\n|\n\s*={3,}\s*\n', content)
    
    for section in sections:
        lines = section.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            
            # Look for review fields
            if line.startswith('**ID:**'):
                # Save previous review if exists
                if current_review:
                    reviews.append(current_review)
                current_review = {'id': line.replace('**ID:**', '').strip()}
                
            elif line.startswith('**Category:**') and current_review:
                current_review['category'] = line.replace('**Category:**', '').strip()
                
            elif line.startswith('**Requirement:**') and current_review:
                current_review['requirement'] = line.replace('**Requirement:**', '').strip()
                
            elif line.startswith('**Description:**') and current_review:
                current_review['description'] = line.replace('**Description:**', '').strip()
                
            elif line.startswith('**ISO Clause:**') and current_review:
                current_review['iso_clause'] = line.replace('**ISO Clause:**', '').strip()
                
            elif line.startswith('**Status:**') and current_review:
                current_review['status'] = line.replace('**Status:**', '').strip()
                
            elif line.startswith('**Comment:**') and current_review:
                current_review['comment'] = line.replace('**Comment:**', '').strip()
                
            elif line.startswith('**Hint for improvement:**') and current_review:
                current_review['hint_for_improvement'] = line.replace('**Hint for improvement:**', '').strip()
    
    # Add last review
    if current_review:
        reviews.append(current_review)
    
    log.info(f"Parsed {len(reviews)} review items from content")
    return reviews

def group_reviews_by_category(reviews, plugin_folder):
    """
    Group review items by their ISO 26262 categories.
    
    Args:
        reviews (list): List of review items
        plugin_folder (str): Path to plugin folder (unused, kept for compatibility)
        
    Returns:
        dict: Reviews grouped by category in logical order
    """
    # Define the logical order of categories
    category_order = [
        "Identification and Classification",
        "Functional Description", 
        "Safety-Related Attributes",
        "Dependencies and Interactions",
        "System Boundaries and Context",
        "Review and Approval",
        "General Requirements"
    ]
    
    # Initialize ordered dictionary with empty lists
    categorized = {category: [] for category in category_order}
    
    for review in reviews:
        # Get category directly from LLM output
        category = review.get('category', 'General Requirements')
        
        # Handle cases where category might be 'N/A' or empty
        if not category or category in ['N/A', '']:
            category = 'General Requirements'
        
        # Ensure the category exists in our dictionary
        if category not in categorized:
            categorized[category] = []
        
        categorized[category].append(review)
    
    # Remove empty categories from the final result
    return {k: v for k, v in categorized.items() if v}

def create_custom_styles(doc, style_prefix="Custom"):
    """
    Create custom styles for Word documents.
    
    Args:
        doc: python-docx Document object
        style_prefix (str): Prefix for style names
    """
    try:
        # Title style
        title_style = doc.styles.add_style(f"{style_prefix}Title", 1)
        title_style.base_style = doc.styles["Normal"]
        title_style.font.name = "Calibri"
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(54, 95, 145)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        # Subtitle style
        subtitle_style = doc.styles.add_style(f"{style_prefix}Subtitle", 1)
        subtitle_style.base_style = doc.styles["Normal"]
        subtitle_style.font.name = "Calibri"
        subtitle_style.font.size = Pt(16)
        subtitle_style.font.color.rgb = RGBColor(54, 95, 145)
        subtitle_style.font.italic = True
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(12)

        # Header style for sections
        header_style = doc.styles.add_style(f"{style_prefix}Header", 1)
        header_style.base_style = doc.styles["Normal"]
        header_style.font.name = "Calibri"
        header_style.font.size = Pt(14)
        header_style.font.bold = True
        header_style.font.color.rgb = RGBColor(54, 95, 145)
        header_style.paragraph_format.space_before = Pt(12)
        header_style.paragraph_format.space_after = Pt(6)

        # Body style
        body_style = doc.styles.add_style(f"{style_prefix}Body", 1)
        body_style.base_style = doc.styles["Normal"]
        body_style.font.name = "Calibri"
        body_style.font.size = Pt(11)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15

        # Guidance label style (for templates)
        guidance_style = doc.styles.add_style(f"{style_prefix}GuidanceLabel", 1)
        guidance_style.base_style = doc.styles["Normal"]
        guidance_style.font.name = "Calibri"
        guidance_style.font.size = Pt(11)
        guidance_style.font.bold = True
        guidance_style.font.color.rgb = RGBColor(0, 102, 204)  # Blue
        guidance_style.paragraph_format.space_before = Pt(6)
        guidance_style.paragraph_format.space_after = Pt(3)

    except Exception as e:
        log.warning(f"Some styles may already exist: {e}")

def add_header_footer(doc, plugin_folder, title_text):
    """
    Add corporate header and footer to Word document.
    
    Args:
        doc: python-docx Document object
        plugin_folder (str): Path to plugin folder
        title_text (str): Text to display in header
    """
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    # === HEADER ===
    table = header.add_table(rows=1, cols=2, width=Inches(8))
    table.autofit = False
    cell_logo, cell_title = table.rows[0].cells

    # Logo (left)
    image_path = os.path.join(plugin_folder, "templates", "logo.png")
    if os.path.exists(image_path):
        paragraph = cell_logo.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.5))
        cell_logo.width = Inches(1.5)
    else:
        cell_logo.text = ""

    # Title (right)
    title_para = cell_title.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title_para.add_run(title_text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(54, 95, 145)

    # Light blue line below header
    line_para = header.add_paragraph()
    p = line_para._element
    pPr = p.get_or_add_pPr()
    bottom_border = OxmlElement('w:pBdr')
    pPr.append(bottom_border)
    bottom_line = OxmlElement('w:bottom')
    bottom_line.set(qn('w:val'), 'single')
    bottom_line.set(qn('w:sz'), '10')
    bottom_line.set(qn('w:space'), '1')
    bottom_line.set(qn('w:color'), 'A0C4E8')
    bottom_border.append(bottom_line)

    # === FOOTER ===
    line_para_footer = footer.add_paragraph()
    p_footer = line_para_footer._element
    pPr_footer = p_footer.get_or_add_pPr()
    top_border = OxmlElement('w:pBdr')
    pPr_footer.append(top_border)
    top_line = OxmlElement('w:top')
    top_line.set(qn('w:val'), 'single')
    top_line.set(qn('w:sz'), '6')
    top_line.set(qn('w:space'), '1')
    top_line.set(qn('w:color'), 'D0E3F5')
    top_border.append(top_line)

    # Footer text
    text_para = footer.add_paragraph()
    text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = text_para.add_run("CONFIDENTIAL – © 2025 Kineton. Generated by Kineton FuSa Agent.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

def add_section_explanation(doc, category):
    """
    Add detailed explanation for ISO 26262 categories.
    
    Args:
        doc: python-docx Document object
        category (str): Category name
    """
    explanations = {
        "Identification and Classification": (
            "This section ensures that the item is uniquely identified within the system architecture or documentation "
            "and properly classified (e.g., as hardware, software, or a system function). This supports traceability from "
            "high-level safety goals down to detailed design elements. Clear identification also enables effective configuration "
            "management and version control throughout the development lifecycle. It is a foundational element for ensuring structured functional safety development."
        ),
        "Functional Description": (
            "This section describes the expected behavior of the item under all operating conditions, including normal, degraded, "
            "and fault modes. It includes definitions of interfaces, timing constraints, and performance requirements. A well-defined "
            "functional description is essential for identifying potential failure scenarios and serves as input to hazard analysis. It helps "
            "ensure that all relevant behaviors are considered when deriving safety requirements."
        ),
        "Safety-Related Attributes": (
            "This section captures key safety-related properties such as safety goals, mitigation strategies, diagnostic coverage, "
            "and safe state definitions. These attributes are derived from the Hazard Analysis and Risk Assessment (HARA) and form the basis "
            "of the functional safety concept. They guide the implementation of safety mechanisms and define how the item contributes to overall "
            "system safety. Proper documentation ensures alignment with ISO 26262 expectations for safety integrity."
        ),
        "Dependencies and Interactions": (
            "This section identifies internal and external dependencies, including interactions with other systems, environmental influences, "
            "and user inputs. Understanding these relationships is critical for defining correct assumptions and boundary conditions during development. "
            "It also supports the identification of potential interference or integration risks that could impact safety. Accurate documentation ensures robust "
            "interface management and system integration."
        ),
        "System Boundaries and Context": (
            "This section defines the physical and logical boundaries of the item, along with environmental conditions and design constraints. "
            "It clarifies where the item operates and under what limitations, such as temperature, vibration, or EMC exposure. These details ensure that "
            "the item is developed and validated under realistic assumptions. Defining this context early supports the creation of accurate test plans and operational profiles."
        ),
        "Review and Approval": (
            "This section confirms that a formal review process was followed and that all necessary approvals were obtained before finalizing the item definition. "
            "It verifies that review minutes, action items, and change records are documented and closed. Configuration management practices should also be applied to maintain "
            "document integrity. This ensures process compliance and provides an auditable trail for quality assurance and functional safety governance."
        )
    }

    explanation = explanations.get(category)
    if explanation:
        paragraph = doc.add_paragraph(explanation)
        paragraph.style = 'Normal'
        paragraph.alignment = 0  # Left-aligned
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.line_spacing = 1.2
        log.debug(f"Added explanation for category: {category}")
    else:
        log.warning(f"No explanation found for category: {category}")

def parse_markdown_content(content):
    """
    Parse markdown-style content for Item Definition documents.
    Handles both actual definitions and templates with guidance.
    
    Args:
        content (str): Markdown-style content
        
    Returns:
        list: List of content sections with type and text
    """
    lines = content.split("\n")
    sections = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Main title (single #)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            sections.append({"type": "title", "text": stripped[2:].strip()})
        
        # Heading 1 (##)
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            sections.append({"type": "heading1", "text": stripped[3:].strip()})
        
        # Heading 2 (###)
        elif stripped.startswith("### "):
            sections.append({"type": "heading2", "text": stripped[4:].strip()})
        
        # Clause references (italic text with *)
        elif stripped.startswith("*Clause:") and stripped.endswith("*"):
            sections.append({"type": "clause", "text": stripped[1:-1].strip()})
        
        # Bold labels (entire line is **Label:**)
        elif stripped.startswith("**") and stripped.endswith("**") and ":" in stripped:
            sections.append({"type": "bold_label", "text": stripped})
        
        # Other italic text (single * at start and end, entire line)
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            sections.append({"type": "italic", "text": stripped})
        
        # Horizontal rule
        elif stripped == "---":
            sections.append({"type": "separator", "text": stripped})
        
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            sections.append({"type": "bullet", "text": stripped[2:].strip()})
        
        # Regular text (may contain inline bold/italic)
        else:
            sections.append({"type": "body", "text": stripped})
    
    return sections