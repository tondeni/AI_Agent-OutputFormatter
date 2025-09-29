# item_definition_dev_doc.py - Item Definition Word document formatter
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from cat.log import log
from .utils import add_header_footer, parse_markdown_content

def create_item_definition_styles(doc):
    """Create custom styles specifically for Item Definition documents."""
    try:
        # Title style
        title_style = doc.styles.add_style("ItemDefTitle", 1)
        title_style.base_style = doc.styles["Normal"]
        title_style.font.name = "Calibri"
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(54, 95, 145)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        # Subtitle style
        subtitle_style = doc.styles.add_style("ItemDefSubtitle", 1)
        subtitle_style.base_style = doc.styles["Normal"]
        subtitle_style.font.name = "Calibri"
        subtitle_style.font.size = Pt(16)
        subtitle_style.font.color.rgb = RGBColor(54, 95, 145)
        subtitle_style.font.italic = True
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(12)

        # Header style for sections
        header_style = doc.styles.add_style("ItemDefHeader", 1)
        header_style.base_style = doc.styles["Normal"]
        header_style.font.name = "Calibri"
        header_style.font.size = Pt(14)
        header_style.font.bold = True
        header_style.font.color.rgb = RGBColor(54, 95, 145)
        header_style.paragraph_format.space_before = Pt(12)
        header_style.paragraph_format.space_after = Pt(6)

        # Body style
        body_style = doc.styles.add_style("ItemDefBody", 1)
        body_style.base_style = doc.styles["Normal"]
        body_style.font.name = "Calibri"
        body_style.font.size = Pt(11)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15

        # Guidance label style (for templates)
        guidance_style = doc.styles.add_style("ItemDefGuidanceLabel", 1)
        guidance_style.base_style = doc.styles["Normal"]
        guidance_style.font.name = "Calibri"
        guidance_style.font.size = Pt(11)
        guidance_style.font.bold = True
        guidance_style.font.color.rgb = RGBColor(0, 102, 204)  # Blue
        guidance_style.paragraph_format.space_before = Pt(6)
        guidance_style.paragraph_format.space_after = Pt(3)

    except Exception as e:
        log.warning(f"Some styles may already exist: {e}")

def process_inline_markdown(paragraph, text):
    """
    Process inline markdown formatting (bold and italic) within text.
    
    Args:
        paragraph: python-docx Paragraph object
        text: Text that may contain **bold** and *italic* markdown
    """
    import re
    
    # Pattern to match **bold** or *italic*
    # This regex splits on markdown while capturing the delimiters
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)

def create_item_definition_docx(content, plugin_folder, system_name):
    """
    Create a formatted Word document for Item Definition.
    Handles both actual item definitions and templates with guidance.
    """
    log.info(f"Creating Item Definition document for {system_name}")
    
    # Detect if this is a template
    is_template = "Template" in content[:500] or "[Item Name]" in system_name
    
    doc = Document()
    
    # Create styles
    create_item_definition_styles(doc)
    
    # Add header/footer
    header_text = f"ISO 26262 Part 3 - Item Definition\n{system_name}"
    if is_template:
        header_text += "\nTEMPLATE"
    
    add_header_footer(doc, plugin_folder, header_text)
    
    # Title page
    doc.add_paragraph('ISO 26262 Part 3 - Item Definition', style="ItemDefTitle")
    doc.add_paragraph(f"System: {system_name}", style="ItemDefSubtitle")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                     style="ItemDefSubtitle")
    
    if is_template:
        template_note = doc.add_paragraph("Document Type: TEMPLATE WITH GUIDANCE", 
                                         style="ItemDefSubtitle")
        template_note.runs[0].font.bold = True
        template_note.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # Parse markdown content into structured sections
    sections = parse_markdown_content(content)
    
    # Render each section with appropriate styling
    for section in sections:
        section_type = section["type"]
        text = section["text"]
        
        if section_type == "title":
            continue
        
        elif section_type == "heading1":
            p = doc.add_paragraph(style="ItemDefHeader")
            process_inline_markdown(p, text)
        
        elif section_type == "heading2":
            p = doc.add_paragraph(style="ItemDefHeader")
            process_inline_markdown(p, text)
            p.runs[0].font.size = Pt(12)
        
        elif section_type == "clause":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(54, 95, 145)
        
        elif section_type == "bold_label":
            # Bold section labels (Guidance, Format, Examples, etc.)
            doc.add_paragraph(text.strip("*"), style="ItemDefGuidanceLabel")
        
        elif section_type == "italic":
            # Italic notes
            p = doc.add_paragraph(text.strip("*"))
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        elif section_type == "separator":
            doc.add_paragraph()
        
        elif section_type == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            process_inline_markdown(p, text)
        
        else:
            # Body text - check if it contains placeholders
            if "[" in text and "]" in text:
                p = doc.add_paragraph(style="ItemDefBody")
                process_inline_markdown(p, text)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    run.italic = True
            else:
                p = doc.add_paragraph(style="ItemDefBody")
                process_inline_markdown(p, text)
    
    log.info("Item Definition document created successfully")
    return doc